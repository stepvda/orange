"""One emitter per output format. Each walks a `blocks.Document` or `blocks.Deck`.

Five formats, and the choice belongs to the reader rather than to this code:
a battlecard is a PDF in a car park, a Word file on a bid manager's desk and an
ODF file on an estate that standardised on LibreOffice, and it is the same
battlecard in all three. That is the whole reason the documents are described in
`documents`/`decks` rather than rendered there.

WHAT EACH FORMAT IS FOR, and what it costs:

  pdf    reportlab. Vector charts, exact geometry, no browser (NFR-05). The
         format to send.
  docx   python-docx. Native styles and tables; charts arrive as high-DPI
         images because Word has no drawing model this code can target.
  odt    odfpy. The same, for the LibreOffice estate.
  pptx   python-pptx. Charts are NATIVE SHAPES — the architect can move a box
         rather than redraw the slide. The format to edit.
  odp    odfpy. Slides with rasterised charts.

The raster fallback is the right way round: the format people send (PDF) and the
format people edit (PPTX) both get true vector output, and the fallback only
applies where a chart is being read rather than worked on.
"""

from __future__ import annotations

import datetime as dt
import io
from pathlib import Path
from typing import Any

from . import blocks as B
from .context import TopicContext

ORANGE_HEX = "F16E00"
INK_HEX = "141414"
MUTED_HEX = "666666"
WARN_HEX = "8A5A00"


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def document_to_pdf(doc: B.Document, ctx: TopicContext, path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import (BaseDocTemplate, Frame, ListFlowable, ListItem, PageBreak,
                                    PageTemplate, Paragraph, Spacer, Table, TableStyle)

    from ..brief import INK, MUTED, ORANGE, RULE, SURFACE, _big, _clip, _styles, _table_style, _text
    from . import charts

    style = _styles()
    width = B.USABLE_WIDTH
    story: list[Any] = []

    for block in doc.blocks:
        if isinstance(block, B.Kicker):
            story.append(Paragraph(_text(block.text), style["kicker"]))
        elif isinstance(block, B.Heading):
            key = {1: "title", 2: "h2", 3: "h3"}.get(block.level, "h2")
            story.append(Paragraph(_text(block.text), style[key]))
        elif isinstance(block, B.Para):
            text = block.text if block.markup else _text(block.text)
            if block.bold and not block.markup:
                text = f"<b>{text}</b>"
            story.append(Paragraph(text, style["small" if block.small else "body"]))
        elif isinstance(block, B.Bullets):
            story.append(ListFlowable(
                [ListItem(Paragraph(item, style["body"]), leftIndent=10) for item in block.items],
                bulletType="bullet", bulletFontSize=5.5, leftIndent=10, bulletOffsetY=-1))
        elif isinstance(block, B.Callout):
            story.append(Paragraph(_text(block.text), style["warn"]))
            story.append(Spacer(1, 2 * mm))
        elif isinstance(block, B.KPIs):
            cells = []
            for index, (value, caption) in enumerate(block.cells):
                tone = block.tones[index] if index < len(block.tones) else ""
                colour = (ORANGE if tone == "accent"
                          else charts.STATUS.get(tone, INK) if tone else INK)
                cells.append(_big(value, caption, colour))
            table = Table([cells], colWidths=[width / max(len(cells), 1)] * len(cells),
                          hAlign="LEFT")
            table.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (0, 0), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6)]))
            story.append(table)
        elif isinstance(block, B.Table):
            rows = []
            if block.headers:
                rows.append([Paragraph(f"<b>{_text(h)}</b>", style["cellhead"])
                             for h in block.headers])
            rows += [[Paragraph(_text(cell), style["cell"]) for cell in row] for row in block.rows]
            widths = block.widths or [1] * len(rows[0])
            total = sum(widths) or 1
            table = Table(rows, colWidths=[width * w / total for w in widths], hAlign="LEFT")
            if block.headers:
                table.setStyle(_table_style())
            else:
                table.setStyle(TableStyle([
                    ("BOX", (0, 0), (-1, -1), 0.4, RULE),
                    ("INNERGRID", (0, 0), (-1, -1), 0.3, RULE),
                    ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, SURFACE]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3)]))
            story.append(table)
        elif isinstance(block, B.Chart):
            story.append(Spacer(1, 2 * mm))
            story.append(block.build(width))
            if block.caption:
                story.append(Paragraph(_text(block.caption), style["small"]))
        elif isinstance(block, B.PageBreak):
            story.append(PageBreak())
        elif isinstance(block, B.Spacer):
            story.append(Spacer(1, block.mm * mm))

    template = BaseDocTemplate(
        str(path), pagesize=A4,
        leftMargin=17 * mm, rightMargin=17 * mm, topMargin=17 * mm, bottomMargin=16 * mm,
        title=f"{ctx.topic_id} — {doc.title}", author="Orange Business Innovation Radar",
        subject=f"Pre-sales collateral for {ctx.statement[:120]}")
    frame = Frame(template.leftMargin, template.bottomMargin, template.width, template.height,
                  id="body")
    footer = f"{ctx.topic_id} · {_clip(ctx.statement, 78)} · {doc.subject}"

    def decorate(canvas, document):
        canvas.saveState()
        canvas.setFillColor(ORANGE)
        canvas.rect(0, A4[1] - 6, A4[0], 6, stroke=0, fill=1)
        canvas.setFont("Helvetica", 6.6)
        canvas.setFillColor(MUTED)
        canvas.drawString(17 * mm, 10 * mm, charts.safe(footer))
        canvas.drawRightString(A4[0] - 17 * mm, 10 * mm, f"{document.page}")
        canvas.restoreState()

    template.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=decorate)])
    template.build(story)


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------

def document_to_docx(doc: B.Document, ctx: TopicContext, path: Path) -> None:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    document = DocxDocument()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    def para(text: str, *, size=10.5, bold=False, italic=False, colour=INK_HEX, space_after=6):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(space_after)
        run = paragraph.add_run(B.plain(text))
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor.from_string(colour)
        return paragraph

    for block in doc.blocks:
        if isinstance(block, B.Kicker):
            para(block.text, size=9, bold=True, colour=ORANGE_HEX, space_after=2)
        elif isinstance(block, B.Heading):
            size = {1: 20, 2: 13, 3: 10.5}.get(block.level, 13)
            para(block.text, size=size, bold=True,
                 colour=MUTED_HEX if block.level == 3 else INK_HEX, space_after=4)
        elif isinstance(block, B.Para):
            para(block.text, size=8.5 if block.small else 10.5, bold=block.bold,
                 colour=MUTED_HEX if block.small else INK_HEX)
        elif isinstance(block, B.Bullets):
            for item in block.items:
                paragraph = document.add_paragraph(B.plain(item), style="List Bullet")
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
        elif isinstance(block, B.Callout):
            para(block.text, size=9.5, italic=True, colour=WARN_HEX)
        elif isinstance(block, B.KPIs):
            table = document.add_table(rows=2, cols=len(block.cells))
            for index, (value, caption) in enumerate(block.cells):
                top = table.cell(0, index).paragraphs[0]
                run = top.add_run(value)
                run.font.size = Pt(16)
                run.font.bold = True
                bottom = table.cell(1, index).paragraphs[0]
                run = bottom.add_run(caption)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor.from_string(MUTED_HEX)
            document.add_paragraph()
        elif isinstance(block, B.Table):
            columns = len(block.headers or (block.rows[0] if block.rows else [""]))
            table = document.add_table(rows=0, cols=columns)
            table.style = "Table Grid"
            if block.headers:
                cells = table.add_row().cells
                for index, header in enumerate(block.headers):
                    run = cells[index].paragraphs[0].add_run(header)
                    run.font.bold = True
                    run.font.size = Pt(9)
            for row in block.rows:
                cells = table.add_row().cells
                for index, value in enumerate(row[:columns]):
                    run = cells[index].paragraphs[0].add_run(B.plain(value))
                    run.font.size = Pt(9)
            document.add_paragraph()
        elif isinstance(block, B.Chart):
            # Word has no drawing model this code can target, so the same
            # reportlab geometry is rasterised at print resolution. It reads
            # identically to the PDF; it just cannot be edited in place.
            width_mm, height_mm = block.size_mm()
            document.add_picture(io.BytesIO(block.png()), width=Inches(width_mm / 25.4))
            if block.caption:
                para(block.caption, size=8.5, colour=MUTED_HEX)
        elif isinstance(block, B.PageBreak):
            document.add_page_break()
        elif isinstance(block, B.Spacer):
            document.add_paragraph()

    document.core_properties.title = f"{ctx.topic_id} — {doc.title}"
    document.core_properties.author = "Orange Business Innovation Radar"
    document.core_properties.subject = ctx.statement[:200]
    document.save(str(path))


# ---------------------------------------------------------------------------
# ODF text
# ---------------------------------------------------------------------------

def _odf_styles(document: Any) -> dict[str, Any]:
    """The house style, declared once as real ODF automatic styles.

    Written as styles rather than as per-run formatting so the file opens in
    LibreOffice as something a person can restyle, which is the only reason to
    ship ODF rather than a PDF in the first place.
    """
    from odf.style import (ParagraphProperties, Style, TableColumnProperties,
                           TableCellProperties, TextProperties)

    made: dict[str, Any] = {}
    spec = {
        "Kicker": dict(size="9pt", weight="bold", colour=f"#{ORANGE_HEX}", after="0.05cm"),
        "Title1": dict(size="20pt", weight="bold", colour=f"#{INK_HEX}", after="0.15cm"),
        "Title2": dict(size="13pt", weight="bold", colour=f"#{INK_HEX}",
                       before="0.35cm", after="0.12cm"),
        "Title3": dict(size="10.5pt", weight="bold", colour=f"#{MUTED_HEX}",
                       before="0.25cm", after="0.08cm"),
        "Body": dict(size="10.5pt", colour=f"#{INK_HEX}", after="0.16cm"),
        "Small": dict(size="8.5pt", colour=f"#{MUTED_HEX}", after="0.12cm"),
        "BodyBold": dict(size="10.5pt", weight="bold", colour=f"#{INK_HEX}", after="0.16cm"),
        "Warn": dict(size="9.5pt", style="italic", colour=f"#{WARN_HEX}", after="0.2cm"),
        "KpiValue": dict(size="16pt", weight="bold", colour=f"#{INK_HEX}", after="0cm"),
        "KpiCaption": dict(size="8pt", colour=f"#{MUTED_HEX}", after="0.2cm"),
        "CellHead": dict(size="9pt", weight="bold", colour=f"#{INK_HEX}", after="0cm"),
        "Cell": dict(size="9pt", colour=f"#{INK_HEX}", after="0cm"),
    }
    for name, conf in spec.items():
        style = Style(name=name, family="paragraph")
        style.addElement(TextProperties(
            fontsize=conf["size"], fontweight=conf.get("weight", "normal"),
            fontstyle=conf.get("style", "normal"), color=conf["colour"]))
        style.addElement(ParagraphProperties(
            margintop=conf.get("before", "0cm"), marginbottom=conf.get("after", "0.15cm")))
        document.automaticstyles.addElement(style)
        made[name] = style

    cell = Style(name="CellBox", family="table-cell")
    cell.addElement(TableCellProperties(border="0.5pt solid #DDDDDD", padding="0.08cm"))
    document.automaticstyles.addElement(cell)
    made["CellBox"] = cell

    graphic = Style(name="ChartFrame", family="graphic")
    document.automaticstyles.addElement(graphic)
    made["ChartFrame"] = graphic
    return made


def document_to_odt(doc: B.Document, ctx: TopicContext, path: Path) -> None:
    from odf.draw import Frame, Image
    from odf.opendocument import OpenDocumentText
    from odf.table import Table as OdfTable, TableCell, TableColumn, TableRow
    from odf.text import H, P, Span

    document = OpenDocumentText()
    styles = _odf_styles(document)
    body = document.text

    def para(text: str, style_name: str) -> None:
        body.addElement(P(stylename=styles[style_name], text=B.plain(text)))

    for block in doc.blocks:
        if isinstance(block, B.Kicker):
            para(block.text, "Kicker")
        elif isinstance(block, B.Heading):
            body.addElement(H(outlinelevel=block.level,
                              stylename=styles[f"Title{min(block.level, 3)}"],
                              text=B.plain(block.text)))
        elif isinstance(block, B.Para):
            para(block.text, "Small" if block.small else "BodyBold" if block.bold else "Body")
        elif isinstance(block, B.Bullets):
            for item in block.items:
                para(f"•  {B.plain(item)}", "Body")
        elif isinstance(block, B.Callout):
            para(block.text, "Warn")
        elif isinstance(block, B.KPIs):
            for value, caption in block.cells:
                para(value, "KpiValue")
                para(caption, "KpiCaption")
        elif isinstance(block, B.Table):
            table = OdfTable(name=f"t{id(block)}")
            columns = len(block.headers or (block.rows[0] if block.rows else [""]))
            table.addElement(TableColumn(numbercolumnsrepeated=columns))
            source = ([block.headers] if block.headers else []) + block.rows
            for index, row in enumerate(source):
                odf_row = TableRow()
                for value in row[:columns]:
                    odf_cell = TableCell(valuetype="string", stylename=styles["CellBox"])
                    odf_cell.addElement(P(
                        stylename=styles["CellHead" if index == 0 and block.headers else "Cell"],
                        text=B.plain(value)))
                    odf_row.addElement(odf_cell)
                table.addElement(odf_row)
            body.addElement(table)
            para("", "Small")
        elif isinstance(block, B.Chart):
            width_mm, height_mm = block.size_mm()
            href = document.addPicture(f"chart{id(block)}.png", "image/png", block.png())
            frame = Frame(stylename=styles["ChartFrame"], anchortype="paragraph",
                          width=f"{width_mm:.1f}mm", height=f"{height_mm:.1f}mm")
            frame.addElement(Image(href=href))
            wrapper = P()
            wrapper.addElement(frame)
            body.addElement(wrapper)
            if block.caption:
                para(block.caption, "Small")
        elif isinstance(block, B.Spacer):
            para("", "Small")
        elif isinstance(block, B.PageBreak):
            # ODF page breaks need a dedicated paragraph style; a blank line is
            # honest and does not risk an invalid document.
            para("", "Body")

    document.save(str(path))


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------

def document_to_md(doc: B.Document, ctx: TopicContext, path: Path) -> None:
    """Markdown, for the artefacts that get pasted rather than opened.

    Charts become a named placeholder rather than a broken image link: a
    Markdown file has nowhere to put a binary, and a dead `![](chart.png)` in
    somebody's mail client is worse than a line saying what is missing and
    where to get it.
    """
    lines: list[str] = []
    for block in doc.blocks:
        if isinstance(block, B.Kicker):
            lines += [f"`{B.plain(block.text)}`", ""]
        elif isinstance(block, B.Heading):
            lines += ["#" * max(1, min(block.level, 4)) + " " + B.plain(block.text), ""]
        elif isinstance(block, B.Para):
            text = B.plain(block.text)
            lines += [f"_{text}_" if block.small else f"**{text}**" if block.bold else text, ""]
        elif isinstance(block, B.Bullets):
            lines += [f"- {B.plain(item)}" for item in block.items] + [""]
        elif isinstance(block, B.Callout):
            lines += [f"> **Note:** {B.plain(block.text)}", ""]
        elif isinstance(block, B.KPIs):
            lines += [" · ".join(f"**{v}** {c}" for v, c in block.cells), ""]
        elif isinstance(block, B.Table):
            headers = block.headers or [""] * len(block.rows[0] if block.rows else [""])
            lines += ["| " + " | ".join(headers) + " |",
                      "|" + "|".join("---" for _ in headers) + "|"]
            lines += ["| " + " | ".join(B.plain(c).replace("|", "\\|") for c in row) + " |"
                      for row in block.rows]
            lines += [""]
        elif isinstance(block, B.Chart):
            lines += [f"_[chart: {block.caption or 'see the PDF or PowerPoint version'}]_", ""]
        elif isinstance(block, B.PageBreak):
            lines += ["---", ""]
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


# ---------------------------------------------------------------------------
# ODF presentation
# ---------------------------------------------------------------------------

def deck_to_odp(deck: B.Deck, ctx: TopicContext, path: Path) -> None:
    from odf.draw import Frame, Image, Page, TextBox
    from odf.opendocument import OpenDocumentPresentation
    from odf.presentation import Notes
    from odf.style import (DrawingPageProperties, GraphicProperties, MasterPage,
                           PageLayout, PageLayoutProperties, ParagraphProperties, Style,
                           TextProperties)
    from odf.text import P

    document = OpenDocumentPresentation()
    layout = PageLayout(name="DeckLayout")
    document.automaticstyles.addElement(layout)
    layout.addElement(PageLayoutProperties(
        margin="0cm", pagewidth="33.87cm", pageheight="19.05cm", printorientation="landscape"))
    master = MasterPage(name="Deck", pagelayoutname=layout)
    document.masterstyles.addElement(master)

    def text_style(name: str, size: str, weight: str, colour: str,
                   align: str = "left") -> tuple[Any, Any]:
        """A frame style AND the paragraph style that actually sizes the text.

        Two styles, not one, because ODF applies them at different levels: the
        presentation-family style dresses the `draw:frame`, and a renderer takes
        the font from the PARAGRAPH inside it. Setting the size only on the
        frame — which is the obvious thing to do, and what this did first — left
        LibreOffice falling back to its own defaults, so a 24pt title and an 8pt
        footer came out nearly the same size.
        """
        frame_style = Style(name=name, family="presentation")
        frame_style.addElement(GraphicProperties(fill="none", stroke="none"))
        frame_style.addElement(ParagraphProperties(textalign=align))
        frame_style.addElement(TextProperties(fontsize=size, fontweight=weight, color=colour))
        document.automaticstyles.addElement(frame_style)

        paragraph_style = Style(name=f"{name}-p", family="paragraph")
        paragraph_style.addElement(ParagraphProperties(textalign=align))
        paragraph_style.addElement(TextProperties(
            fontsize=size, fontweight=weight, color=colour,
            fontfamily="Arial", fontsizeasian=size, fontsizecomplex=size))
        document.automaticstyles.addElement(paragraph_style)
        return frame_style, paragraph_style

    styles = {
        "cover": text_style("s-cover", "34pt", "bold", f"#{INK_HEX}"),
        "title": text_style("s-title", "24pt", "bold", f"#{INK_HEX}"),
        "sub": text_style("s-sub", "11pt", "normal", f"#{MUTED_HEX}"),
        "bullet": text_style("s-bullet", "14pt", "normal", f"#{INK_HEX}"),
        "kicker": text_style("s-kicker", "12pt", "bold", f"#{ORANGE_HEX}"),
        "footer": text_style("s-footer", "8pt", "normal", f"#{MUTED_HEX}"),
    }
    # A filled band rather than a text frame: ODF has no page-border primitive,
    # and a 1.6mm rectangle in the brand colour is what the PDFs draw anyway.
    rule_style = Style(name="s-rule", family="presentation")
    rule_style.addElement(GraphicProperties(fill="solid", fillcolor=f"#{ORANGE_HEX}",
                                            stroke="none"))
    document.automaticstyles.addElement(rule_style)
    image_style = Style(name="s-image", family="presentation")
    image_style.addElement(GraphicProperties(stroke="none", fill="none"))
    document.automaticstyles.addElement(image_style)

    def box(page: Any, style: tuple[Any, Any], x: float, y: float, w: float, h: float,
            lines: list[str]) -> None:
        frame_style, paragraph_style = style
        frame = Frame(stylename=frame_style, width=f"{w}cm", height=f"{h}cm",
                      x=f"{x}cm", y=f"{y}cm")
        textbox = TextBox()
        for line in lines:
            textbox.addElement(P(stylename=paragraph_style, text=B.plain(line)))
        frame.addElement(textbox)
        page.addElement(frame)

    def chrome(page: Any) -> None:
        """The Orange rule at the head and the space in the footer.

        The same six points of colour every PDF and PowerPoint slide in this
        package carries. Without it an ODF deck is the one artefact in the pack
        that does not look like it came from the same place, which is exactly
        the thing the shared house style exists to prevent.
        """
        rule = Frame(stylename=rule_style, width="33.87cm", height="0.16cm", x="0cm", y="0cm")
        rule.addElement(TextBox())
        page.addElement(rule)
        box(page, styles["footer"], 2, 17.6, 29, 0.8,
            [f"{ctx.topic_id} · {B.plain(ctx.statement)[:96]}"])

    for slide in deck.slides:
        page = Page(masterpagename=master)
        document.presentation.addElement(page)
        chrome(page)

        if slide.kind == "cover":
            box(page, styles["kicker"], 2, 5.2, 29, 1, [f"{ctx.topic_id} · {ctx.triple}"])
            box(page, styles["cover"], 2, 6.4, 29, 3, [slide.title])
            box(page, styles["sub"], 2, 9.8, 29, 2, [ctx.statement, slide.strapline])
        elif slide.kind == "section":
            box(page, styles["title"], 2, 7.5, 29, 2, [slide.title])
            if slide.subtitle:
                box(page, styles["sub"], 2, 9.6, 29, 2, [slide.subtitle])
        else:
            box(page, styles["title"], 2, 1.2, 29, 1.6, [slide.title])
            top = 3.0
            if slide.subtitle:
                box(page, styles["sub"], 2, 2.8, 29, 1, [slide.subtitle])
                top = 4.0
            if slide.bullets:
                box(page, styles["bullet"], 2, top, 29, 1.1 * len(slide.bullets),
                    [f"•  {item}" for item in slide.bullets])
                top += 1.1 * len(slide.bullets) + 0.4
            if slide.rows:
                box(page, styles["bullet"], 2, top, 29, 1.0 * len(slide.rows),
                    [f"{key}:  {value}" for key, value in slide.rows])
            if slide.chart is not None:
                width_mm, height_mm = slide.chart.size_mm()
                # Fit inside what is left of the page rather than at native
                # size: a chart designed for an A4 column is taller than the
                # space under a slide title, and an overflowing image is worse
                # than a smaller one.
                available_h = max(19.05 - top - 1.6, 3.0)
                scale = min(29.0 / (width_mm / 10.0), available_h / (height_mm / 10.0))
                href = document.addPicture(f"c{id(slide)}.png", "image/png", slide.chart.png())
                frame = Frame(stylename=image_style,
                              width=f"{width_mm / 10.0 * scale:.2f}cm",
                              height=f"{height_mm / 10.0 * scale:.2f}cm",
                              x="2cm", y=f"{top:.2f}cm")
                frame.addElement(Image(href=href))
                page.addElement(frame)

        if slide.notes:
            notes = Notes()
            note_frame = Frame(width="20cm", height="10cm", x="2cm", y="2cm")
            textbox = TextBox()
            textbox.addElement(P(text=B.plain(slide.notes)))
            note_frame.addElement(textbox)
            notes.addElement(note_frame)
            page.addElement(notes)

    document.save(str(path))


# ---------------------------------------------------------------------------
# A deck as PDF
# ---------------------------------------------------------------------------

def deck_to_pdf(deck: B.Deck, ctx: TopicContext, path: Path) -> None:
    """Slides on landscape pages, one slide per page.

    Not a document with headings: somebody asking for the deck as PDF wants the
    deck — one idea per page, in the order they will present it — and flowing it
    into A4 portrait would destroy the only property that made it a deck.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas as pdfcanvas

    from ..brief import INK, MUTED, ORANGE, RULE, _clip
    from . import charts

    page_w, page_h = landscape(A4)
    margin = 18 * mm
    body_w = page_w - 2 * margin
    surface = pdfcanvas.Canvas(str(path), pagesize=(page_w, page_h))
    surface.setTitle(f"{ctx.topic_id} — {deck.title}")
    surface.setAuthor("Orange Business Innovation Radar")

    def wrapped(text: str, font: str, size: float, width: float) -> list[str]:
        words, lines, current = charts.safe(text).split(), [], ""
        for word in words:
            trial = f"{current} {word}".strip()
            if surface.stringWidth(trial, font, size) <= width:
                current = trial
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
        return lines

    for slide in deck.slides:
        surface.setFillColor(ORANGE)
        surface.rect(0, page_h - 5, page_w, 5, stroke=0, fill=1)
        surface.setFillColor(MUTED)
        surface.setFont("Helvetica", 7)
        surface.drawString(margin, 9 * mm, charts.safe(f"{ctx.topic_id} · {_clip(ctx.statement, 96)}"))

        if slide.kind == "cover":
            surface.setFillColor(ORANGE)
            surface.rect(0, page_h * 0.42, 4 * mm, 46 * mm, stroke=0, fill=1)
            surface.setFillColor(MUTED)
            surface.setFont("Helvetica-Bold", 10)
            surface.drawString(margin, page_h * 0.62, charts.safe(f"{ctx.topic_id} · {ctx.triple}"))
            surface.setFillColor(INK)
            surface.setFont("Helvetica-Bold", 28)
            for index, line in enumerate(wrapped(slide.title, "Helvetica-Bold", 28, body_w)[:2]):
                surface.drawString(margin, page_h * 0.55 - index * 34, line)
            surface.setFont("Helvetica", 12)
            for index, line in enumerate(wrapped(ctx.statement, "Helvetica", 12, body_w)[:2]):
                surface.drawString(margin, page_h * 0.36 - index * 16, line)
            surface.setFillColor(MUTED)
            surface.setFont("Helvetica", 9.5)
            for index, line in enumerate(wrapped(slide.strapline, "Helvetica", 9.5, body_w)[:3]):
                surface.drawString(margin, page_h * 0.27 - index * 13, line)
        else:
            cursor = page_h - 26 * mm
            surface.setFillColor(INK)
            surface.setFont("Helvetica-Bold", 19)
            for line in wrapped(slide.title, "Helvetica-Bold", 19, body_w)[:2]:
                surface.drawString(margin, cursor, line)
                cursor -= 24
            if slide.subtitle:
                surface.setFillColor(MUTED)
                surface.setFont("Helvetica", 9.5)
                for line in wrapped(slide.subtitle, "Helvetica", 9.5, body_w)[:2]:
                    surface.drawString(margin, cursor, line)
                    cursor -= 13
            cursor -= 6
            surface.setFillColor(INK)
            surface.setFont("Helvetica", 12.5)
            for bullet in slide.bullets[:6]:
                surface.setFillColor(ORANGE)
                surface.circle(margin + 2, cursor + 4, 1.6, stroke=0, fill=1)
                surface.setFillColor(INK)
                for index, line in enumerate(wrapped(bullet, "Helvetica", 12.5, body_w - 8 * mm)[:2]):
                    surface.drawString(margin + 8 * mm, cursor - index * 15, line)
                    cursor -= 15 if index else 0
                cursor -= 20
            for key, value in slide.rows:
                surface.setFillColor(MUTED)
                surface.setFont("Helvetica-Bold", 10)
                surface.drawString(margin, cursor, charts.safe(key))
                surface.setFillColor(INK)
                surface.setFont("Helvetica", 10)
                surface.drawString(margin + 55 * mm, cursor, charts.safe(value)[:120])
                cursor -= 16

            if slide.chart is not None:
                flowable = slide.chart.build(body_w)
                _, height = flowable.wrapOn(surface, body_w, page_h)
                # Scaled to whatever vertical room the title and bullets left,
                # rather than clipped at the bottom of the page.
                available = cursor - 16 * mm
                scale = min(1.0, available / height) if height else 1.0
                surface.saveState()
                surface.translate(margin, cursor - height * scale)
                surface.scale(scale, scale)
                flowable.drawOn(surface, 0, 0)
                surface.restoreState()

        surface.showPage()

    surface.save()
