#!/usr/bin/env python3
"""Build the Next Steps document.

    python3 docs/generators/build_next_steps.py

Self-contained on purpose: the shared `docs/_build/docx_kit.py` helper that the
Functional Design Document and the Technical Architecture were assembled from is
no longer in the tree, so the house style is carried here rather than imported.
Colours, fonts and block shapes match those documents exactly.

Figures in the text are read from the working database where one is present, so
this document cannot quote a count the radar does not hold. Where the database
is absent the recorded snapshot below is used and the cover says which.
"""
from __future__ import annotations

import pathlib
import sqlite3
import sys

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent.parent
DOCS = ROOT / "docs"
DB = ROOT / "data" / "radar.db"
OUT = DOCS / "Orange_Innovation_Radar_Next_Steps.docx"

# ---------------------------------------------------------------- house style
ORANGE = RGBColor(0xE8, 0x6A, 0x00)
ORANGE_DARK = RGBColor(0xA8, 0x3E, 0x00)
INK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x5A, 0x5A, 0x5A)
GREY_LIGHT = RGBColor(0x8A, 0x8A, 0x8A)
BLUE = RGBColor(0x2F, 0x6F, 0xB0)
GREEN = RGBColor(0x2E, 0x7D, 0x5B)
RED = RGBColor(0xA8, 0x28, 0x20)

SH_ORANGE = "FFF1E3"
SH_GREY = "F2F2F2"
SH_HEADER = "3C3C3C"
SH_BLUE = "E6EFF8"
SH_GREEN = "E4F1EA"
SH_RED = "FBECEA"

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"


def _shade(el, fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), fill)
    el.append(sh)


def _borders(el, sides=("top", "left", "bottom", "right"), sz=4, color="D9D9D9", val="single"):
    borders = OxmlElement("w:pBdr") if el.tag.endswith("}pPr") else OxmlElement("w:tcBorders")
    for s in sides:
        e = OxmlElement(f"w:{s}")
        e.set(qn("w:val"), val)
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    el.append(borders)


def _field(paragraph, instr):
    r = paragraph.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin")
    r._r.append(fc)
    r = paragraph.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    r._r.append(it)
    r = paragraph.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "separate")
    r._r.append(fc)
    paragraph.add_run("…")
    r = paragraph.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "end")
    r._r.append(fc)


class Doc:
    def __init__(self, title: str):
        self.d = Document()
        self.title_text = title
        self._setup_styles()
        self._setup_page(self.d.sections[0])
        self._footer(self.d.sections[0])
        self._landscape = False

    def _setup_page(self, sec, landscape=False):
        if landscape:
            sec.orientation = WD_ORIENT.LANDSCAPE
            sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
            sec.top_margin, sec.bottom_margin = Cm(1.5), Cm(1.5)
            sec.left_margin, sec.right_margin = Cm(1.5), Cm(1.5)
        else:
            sec.orientation = WD_ORIENT.PORTRAIT
            sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
            sec.top_margin, sec.bottom_margin = Cm(2.2), Cm(2.0)
            sec.left_margin, sec.right_margin = Cm(2.2), Cm(2.2)

    def _setup_styles(self):
        st = self.d.styles
        n = st["Normal"]
        n.font.name = BODY_FONT
        n.font.size = Pt(10.5)
        n.font.color.rgb = INK
        n.paragraph_format.space_after = Pt(7)
        n.paragraph_format.line_spacing = 1.16
        n.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), BODY_FONT)
        for name, size, color, before, after, bold in [
            ("Heading 1", 19, ORANGE_DARK, 22, 8, True),
            ("Heading 2", 14, INK, 16, 5, True),
            ("Heading 3", 11.5, ORANGE_DARK, 12, 4, True),
            ("Heading 4", 10.5, GREY, 10, 3, True),
        ]:
            s = st[name]
            s.font.name = BODY_FONT
            s.font.size = Pt(size)
            s.font.color.rgb = color
            s.font.bold = bold
            s.paragraph_format.space_before = Pt(before)
            s.paragraph_format.space_after = Pt(after)
            s.paragraph_format.keep_with_next = True

    def _footer(self, sec):
        p = sec.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(self.title_text + "    ·    page ")
        _field(p, " PAGE ")
        for run in p.runs:
            run.font.size = Pt(8); run.font.color.rgb = GREY_LIGHT; run.font.name = BODY_FONT

    # ------------------------------------------------------------- blocks
    def cover(self, title, subtitle, kicker, meta_rows, statement=None):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(70)
        r = p.add_run(kicker)
        r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = ORANGE; r.font.name = BODY_FONT

        p = self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = INK; r.font.name = BODY_FONT

        p = self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(26)
        r = p.add_run(subtitle)
        r.font.size = Pt(14); r.font.color.rgb = GREY; r.font.name = BODY_FONT

        self.rule(ORANGE)

        if statement:
            p = self.d.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(20)
            r = p.add_run(statement)
            r.font.size = Pt(11); r.font.color.rgb = INK; r.font.italic = True; r.font.name = BODY_FONT

        t = self.d.add_table(rows=0, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.autofit = False
        for k, v in meta_rows:
            c0, c1 = t.add_row().cells
            c0.width = Cm(3.6); c1.width = Cm(13.0)
            r = c0.paragraphs[0].add_run(k)
            r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = GREY; r.font.name = BODY_FONT
            r = c1.paragraphs[0].add_run(v)
            r.font.size = Pt(9); r.font.color.rgb = INK; r.font.name = BODY_FONT
            for c in (c0, c1):
                c.paragraphs[0].paragraph_format.space_after = Pt(2)
                _borders(c._tc.get_or_add_tcPr(), sz=0, val="none")
        self.pagebreak()

    def rule(self, color=ORANGE, sz=12):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        bd = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), str(sz)); b.set(qn("w:space"), "1")
        b.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
        bd.append(b)
        p._p.get_or_add_pPr().append(bd)
        return p

    def toc(self, entries):
        self.h1("Contents")
        for level, text in entries:
            p = self.d.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.0 if level == 1 else 0.8)
            r = p.add_run(text)
            r.font.name = BODY_FONT
            r.font.size = Pt(10.5 if level == 1 else 9.5)
            r.font.bold = level == 1
            r.font.color.rgb = INK if level == 1 else GREY
        self.pagebreak()

    def _heading(self, text, level):
        h = self.d.add_heading(text, level=level)
        for r in h.runs:
            r.font.name = BODY_FONT
        return h

    def h1(self, text):
        return self._heading(text, 1)

    def h2(self, text):
        return self._heading(text, 2)

    def h3(self, text):
        return self._heading(text, 3)

    def p(self, text, size=10.5, color=INK, italic=False, bold=False, after=7, before=0, indent=0):
        par = self.d.add_paragraph()
        par.paragraph_format.space_after = Pt(after)
        par.paragraph_format.space_before = Pt(before)
        if indent:
            par.paragraph_format.left_indent = Cm(indent)
        self._rich(par, text, size, color, italic, bold)
        return par

    def _rich(self, par, text, size, color, italic, bold):
        """**bold**, *italic* and `mono` inline markers.

        Bold is matched before italic so a `**` opener is never read as a lone
        `*` — the alternation order is the whole of that guarantee.
        """
        import re
        for chunk in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text):
            if not chunk:
                continue
            r = par.add_run()
            chunk_italic = italic
            if chunk.startswith("**") and chunk.endswith("**"):
                r.text = chunk[2:-2]; r.font.bold = True; r.font.name = BODY_FONT
            elif chunk.startswith("`") and chunk.endswith("`"):
                r.text = chunk[1:-1]; r.font.name = MONO_FONT; r.font.size = Pt(size - 1.0)
                r.font.color.rgb = ORANGE_DARK
            elif len(chunk) > 2 and chunk.startswith("*") and chunk.endswith("*"):
                r.text = chunk[1:-1]; r.font.name = BODY_FONT; r.font.bold = bold
                chunk_italic = True
            else:
                r.text = chunk; r.font.name = BODY_FONT; r.font.bold = bold
            if r.font.size is None:
                r.font.size = Pt(size)
            if r.font.color.rgb is None:
                r.font.color.rgb = color
            r.font.italic = chunk_italic

    def bullets(self, items, size=10.5, style="List Bullet", after=3):
        for it in items:
            par = self.d.add_paragraph(style=style)
            par.paragraph_format.space_after = Pt(after)
            par.paragraph_format.left_indent = Cm(0.7)
            self._rich(par, it, size, INK, False, False)

    def numbers(self, items, size=10.5):
        self.bullets(items, size=size, style="List Number")

    def callout(self, title, body, fill=SH_ORANGE, accent=ORANGE_DARK):
        lines = body if isinstance(body, list) else [body]
        t = self.d.add_table(rows=1, cols=1)
        cell = t.cell(0, 0)
        _shade(cell._tc.get_or_add_tcPr(), fill)
        _borders(cell._tc.get_or_add_tcPr(), ("top", "bottom", "right"), sz=0, val="none")
        _borders(cell._tc.get_or_add_tcPr(), ("left",), sz=18,
                 color="%02X%02X%02X" % (accent[0], accent[1], accent[2]))
        cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
        if title:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(title)
            r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = accent; r.font.name = BODY_FONT
        for i, line in enumerate(lines):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2 if i < len(lines) - 1 else 0)
            self._rich(p, line, 9.5, INK, False, False)
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)

    def table(self, headers, rows, widths=None, size=9, header_fill=SH_HEADER, zebra=True, first_bold=False):
        t = self.d.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            _shade(c._tc.get_or_add_tcPr(), header_fill)
            c.paragraphs[0].paragraph_format.space_after = Pt(1)
            c.paragraphs[0].paragraph_format.space_before = Pt(1)
            r = c.paragraphs[0].add_run(h)
            r.font.size = Pt(size); r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.name = BODY_FONT
        for j, row in enumerate(rows):
            cells = t.add_row().cells
            for i, val in enumerate(row):
                c = cells[i]
                if zebra and j % 2 == 1:
                    _shade(c._tc.get_or_add_tcPr(), SH_GREY)
                pr = c.paragraphs[0]
                pr.paragraph_format.space_after = Pt(1)
                pr.paragraph_format.space_before = Pt(1)
                self._rich(pr, str(val), size, INK, False, first_bold and i == 0)
        for row in t.rows:
            for c in row.cells:
                _borders(c._tc.get_or_add_tcPr(), sz=2, color="C8C8C8")
        if widths:
            t.autofit = False
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Cm(w)
        self.d.add_paragraph().paragraph_format.space_after = Pt(3)
        return t

    def pagebreak(self):
        self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def save(self, path):
        self.d.save(path)
        print("wrote", path)


# ------------------------------------------------------------------- figures
#: Snapshot recorded when this document was written, used only when no working
#: database is present. The radar's own rule applies to its roadmap: a count
#: that cannot be read is a count that says where it came from.
SNAPSHOT = {
    "signals": 11498,
    "raw_items": 11488,
    "spaces": 456,
    "links": 5311,
    "links_confirmed": 0,
    "links_rejected": 0,
    "links_l0": 146,
    "sizes": 762,
    "sized_bottom_up": 340,
    "competition": 220,
    "nodes": 181,
    "edges": 182,
    "refreshes": 59,
    "feedback": 1,
    "internal": 10,
    "assessments": 10,
    "pattern_decisions": 0,
    "observations": 56385,
    "series": 5,
    "source_ids": 34,
    "competitor_pages": 1745,
    "competitor_profiles": 65,
    "competitor_unprofiled": 12,
    "shortlisted": 453,
}


def figures():
    """Read the counts from the working database, falling back to SNAPSHOT."""
    if not DB.exists():
        return dict(SNAPSHOT), False
    f = dict(SNAPSHOT)
    try:
        c = sqlite3.connect(f"file:{DB}?mode=ro", uri=True)
        one = lambda q: c.execute(q).fetchone()[0]
        f.update(
            signals=one("select count(*) from signals"),
            raw_items=one("select count(*) from raw_items"),
            spaces=one("select count(*) from opportunity_spaces"),
            links=one("select count(*) from opportunity_links"),
            links_confirmed=one("select count(*) from opportunity_links where confirmed_by is not null"),
            links_rejected=one("select count(*) from opportunity_links where rejected=1"),
            links_l0=one("select count(*) from opportunity_links where link_type='L0'"),
            sizes=one("select count(*) from market_sizes"),
            sized_bottom_up=one("select count(distinct opportunity_id) from market_sizes"
                                " where method='bottom_up_adoption'"),
            competition=one("select count(distinct opportunity_id) from topic_competition"),
            nodes=one("select count(*) from graph_nodes"),
            edges=one("select count(*) from graph_edges"),
            refreshes=one("select count(*) from refreshes"),
            feedback=one("select count(*) from feedback"),
            internal=one("select count(*) from internal_signals"),
            assessments=one("select count(*) from assessments"),
            pattern_decisions=one("select count(*) from link_pattern_decisions"),
            observations=one("select count(*) from reference_observations"),
            series=one("select count(*) from reference_series"),
            source_ids=one("select count(distinct source_id) from signals"),
            competitor_pages=one("select count(*) from competitor_pages"),
            competitor_profiles=one("select count(*) from competitor_profiles"),
            competitor_unprofiled=one("select count(*) from competitor_profiles"
                                      " where status <> 'profiled'"),
            shortlisted=one("select count(*) from workflow_state where stage='shortlisted'"),
        )
        c.close()
        return f, True
    except sqlite3.Error as exc:                     # a locked or partial file
        print(f"  database unreadable ({exc}); using the recorded snapshot", file=sys.stderr)
        return dict(SNAPSHOT), False


def sources_summary():
    """Read the source catalogue and the connector registry."""
    out = {"total": 42, "enabled": 33, "dark": 9, "connectors": 19, "pending": 18}
    try:
        import yaml
        cat = yaml.safe_load((ROOT / "config" / "sources.yaml").read_text())["sources"]
        enabled = [s for s in cat if s.get("enabled")]
        out["total"] = len(cat)
        out["enabled"] = len(enabled)
        out["dark"] = len(cat) - len(enabled)
        out["pending"] = sum(1 for s in enabled if s.get("terms_checked") == "pending")
        sys.path.insert(0, str(ROOT / "src"))
        from radar.connectors import REGISTRY          # noqa: PLC0415 — optional
        out["connectors"] = len(REGISTRY)
    except Exception as exc:                           # config or import problem
        print(f"  source catalogue unreadable ({exc}); using recorded figures", file=sys.stderr)
    return out


# ------------------------------------------------------------------- content
def main() -> int:
    F, live = figures()
    S = sources_summary()
    n = lambda k: f"{F[k]:,}"

    d = Doc("Orange Innovation Radar — Next Steps")

    d.cover(
        title="Next Steps",
        subtitle="Where the Innovation Radar goes from here — functionally and technically",
        kicker="ORANGE INNOVATION RADAR",
        statement=(
            "The radar closes on two lists: six things deliberately not built, each with its reason, "
            "and six things that need a decision from Orange. This document takes both forward, adds "
            "the directions the machine that now exists makes cheap, and says in each case what the "
            "work is, what it changes, and what it depends on."
        ),
        meta_rows=[
            ("Document", "Next Steps — functional and technical roadmap"),
            ("Version", "1.0"),
            ("Status", "For discussion with Orange Business"),
            ("Audience", "Product owner, architecture, data governance, sales leadership"),
            ("Basis", "The delivered MVP, its source catalogue, its business graph and its working database"),
            ("Figures", ("read from the working database at build time" if live
                         else "recorded snapshot — no working database at build time")),
            ("Companions", "Functional Design Document · Technical Architecture · DECISIONS.md · "
                           "SCORING_FORMULAS.md · STRATEGY_ENGINE_PROPOSAL.md"),
        ],
    )

    d.toc([
        (1, "1 · What this document is"),
        (1, "2 · The four seams the radar extends through"),
        (1, "3 · Replacing and adding data sources"),
        (2, "3.1 Three cases, and only one of them is engineering"),
        (2, "3.2 The connector contract"),
        (2, "3.3 What comes free, and the two things that do not"),
        (2, "3.4 The sources that are catalogued and dark"),
        (1, "4 · Connecting CRM"),
        (2, "4.1 What CRM actually changes, component by component"),
        (2, "4.2 Three integration patterns"),
        (2, "4.3 The crosswalk is the work"),
        (2, "4.4 The wall that has to stay up"),
        (1, "5 · Closing the rest of the not-built list"),
        (2, "5.1 Calibration and backtesting · 5.2 Learned per-role ranking"),
        (2, "5.3 Patents, per-role authorisation, rate limiting, ROI"),
        (1, "6 · The decisions Orange owes the project"),
        (1, "7 · New directions worth building"),
        (2, "7A Refresh diff · 7B Account matching · 7C Ask the evidence"),
        (2, "7D Regulatory deadline calendar · 7E Partner view · 7F Competitor change feed"),
        (2, "7G Collateral into the CRM · 7H Per-country instances · 7I Internal intake"),
        (1, "8 · Platform, operations and technical debt"),
        (2, "8.1 Storage · 8.2 Silent-zero detection · 8.3 Cost per refresh"),
        (2, "8.4 The sovereign path · 8.5 The documentation build · 8.6 Personal data"),
        (1, "9 · A suggested sequence"),
        (1, "10 · What should stay unbuilt"),
    ])

    # ----------------------------------------------------------------- 1
    d.h1("1 · What this document is")
    d.p(
        "The final slide of the radar deck names two lists. On the left, what was deliberately not "
        "built, each row carrying the reason rather than an apology. On the right, what needs a "
        "decision from Orange. This document takes each item forward — what the work actually is, "
        "what changes in the product when it lands, and what it depends on — and then adds a third "
        "list the slide does not have: things nobody has asked for, which the machine that now exists "
        "makes unusually cheap."
    )
    d.p(
        "The ordering principle throughout is **what it changes about what the radar is worth**, not "
        "how hard it is. That ordering is uncomfortable, because it puts governance and procurement "
        "actions with no engineering content at the top, and several genuinely interesting pieces of "
        "modelling near the bottom. Two of the highest-value items in this document are somebody "
        "agreeing to own something. Two of the most expensive-sounding are configuration."
    )
    d.p(
        "One property is treated as non-negotiable in every section: **the radar cannot state a number "
        "it cannot source, and cannot name a thing it did not read.** Everything proposed below is "
        "scoped so that survives it. Section 10 lists the steps that would break it and are therefore "
        "not proposed, including the two that will be the most tempting once CRM data is connected."
    )
    d.callout(
        "Figures in this document",
        [
            f"Counts are {'read from the working database at build time' if live else 'from the snapshot recorded in the generator'}: "
            f"{n('signals')} signals across {n('source_ids')} live sources, {n('spaces')} opportunity spaces, "
            f"{n('links')} links of which **{n('links_confirmed')} are confirmed**, {n('sizes')} market-size "
            f"computations, {n('nodes')} business-graph nodes and {n('refreshes')} recorded refreshes.",
            "They will drift with the next refresh. Where a figure carries an argument in this document, "
            "the argument is about its order of magnitude or about it being zero — not about its "
            "last digit.",
        ],
    )

    # ----------------------------------------------------------------- 2
    d.h1("2 · The four seams the radar extends through")
    d.p(
        "Most of this roadmap is affordable because the radar has exactly four places where new "
        "material enters, and everything downstream of them is indifferent to where the material came "
        "from. Knowing the four is what separates a two-week change from a two-quarter one, so they "
        "are worth stating before any proposal is made."
    )
    d.table(
        ["Seam", "What enters", "Where it lands", "What it moves"],
        [
            ["**The source catalogue**",
             f"`config/sources.yaml` plus the connector registry — {S['total']} sources catalogued, "
             f"{S['enabled']} enabled, {S['connectors']} connector classes implemented",
             "`raw_items` → `signals`",
             "Attractiveness: signal strength, diversity, evidence quality, momentum"],
            ["**The business graph**",
             "`config/business_graph/*.yaml` — offers, partners, references, certifications, analyst "
             "positions, capability pools",
             f"`graph_nodes` / `graph_edges` → `opportunity_links` ({n('nodes')} / {n('edges')} / {n('links')})",
             "Right to win, all seven components; the Planner's capability constraint"],
            ["**Reference data**",
             f"Statistical series on their own cadence — {n('observations')} observations over {n('series')} series",
             "`reference_series` / `reference_observations`",
             "Market size: the denominators, never the signals"],
            ["**The assumption files**",
             "`sizing.yaml`, `economics.yaml`, `settings.yaml`, `source_tiers.yaml`",
             "Version ids stamped onto every artefact they produced",
             "Every published figure, and what may be compared with what"],
        ],
        widths=[3.4, 5.2, 4.0, 4.0],
    )
    d.p(
        "Anything entering through one of those four inherits provenance, replay, version stamping and "
        "the **How was this calculated?** modal at no extra cost. Anything that bypasses them inherits "
        "none of it, and every one of those properties then has to be rebuilt by hand for that one "
        "feature. That is the test to apply to every proposal in this document, and to every proposal "
        "that arrives after it: **which seam does this enter through?** A good answer makes the work "
        "small. No answer is the warning."
    )
    d.p(
        "The reference-data seam is the one most often mistaken for the source seam, and the mistake "
        "is expensive, so it is recorded here. Statistical series are deliberately not ingested as "
        "signals. A signal is a dated event with a publisher, and every attractiveness component "
        "treats it as one — volume, publisher diversity, recency, momentum, tier. An annual "
        "statistical series is none of those things. Pushing tens of thousands of Eurostat cells "
        "through the signal store would corrupt every one of those components while adding nothing to "
        "discovery. The same reasoning will apply, exactly, to CRM data in section 4.",
        after=10,
    )

    # ----------------------------------------------------------------- 3
    d.h1("3 · Replacing and adding data sources")
    d.p(
        "This is the question asked most often about the radar, usually in the form \"what if we don't "
        "want to use that source\" or \"we already licence something better\". The answer is that in "
        "two of the three cases it is a configuration change, and that the interesting part is not the "
        "connector at all — it is what happens to the scores afterwards."
    )

    d.h2("3.1 Three cases, and only one of them is engineering")
    d.table(
        ["Case", "Example", "What it takes", "Effort"],
        [
            ["**A · Same shape, different endpoint**",
             "Replace Google News RSS with a licensed newswire, or an internal press aggregator, that "
             "publishes a feed",
             "Edit the source's `endpoint`, queries, `default_tier` and `terms_checked` in "
             "`config/sources.yaml`. The `rss_feed` / `rss_search` connectors already parse the shape.",
             "Config only — hours"],
            ["**B · New source, existing connector shape**",
             "Any further OCDS procurement portal; any regulator that publishes a feed",
             "Add a catalogue entry pointing at an implemented connector, with its cadence, tier and "
             "geography declared",
             "Config only — hours"],
            ["**C · New connector**",
             "A source with its own API and its own pagination — a national statistics portal, a "
             "licensed analyst feed, an internal data warehouse",
             "Subclass `Connector`, register it with `@register(\"name\")`, implement "
             "`collect(reference_date, since_days)` yielding `CollectedItem`",
             "Typically 80–150 lines"],
        ],
        widths=[3.8, 4.4, 5.4, 3.0],
    )
    d.p(
        f"The catalogue reached {S['total']} sources on {S['connectors']} connector classes precisely "
        "because case B is the common one. Case C is smaller than it sounds, because the framework "
        "already provides per-host rate limiting, conditional requests and cursors, extract "
        "truncation, publisher inference, date parsing and geography normalisation to alpha-2. A "
        "connector implements the part that is genuinely specific to the source and nothing else."
    )

    d.h2("3.2 The connector contract")
    d.p(
        "Four obligations, and they are the whole of it. Three are enforced by the base class; the "
        "first is the connector author's responsibility and is the one that matters.",
        after=5,
    )
    d.bullets([
        "**Never return an item published after the reference date.** The filter is on the "
        "publication date, not the ingestion date. Leakage through late-arriving documents is "
        "invisible unless the pipeline is built from the start to prevent it, and every backtest, "
        "every replay and every momentum slope in the system depends on this one rule holding in "
        "every connector.",
        "**Store by reference.** URL plus a bounded extract, truncated in the base class so no "
        "connector can accidentally mirror a full article. This is a licensing property as much as a "
        "storage one, and it is what makes a source's terms of use a checkable claim rather than a "
        "hope.",
        "**Declare only what the source authoritatively knows.** A procurement notice knows its CPV "
        "codes and its buyer country; an RSS item does not. Fill what is known and leave the rest to "
        "the normalisation stage, which is uniform across sources — hashing, deduplication, "
        "syndication collapse, language detection, tiering.",
        "**Declare geography where the jurisdiction is constant.** A national regulator's items are "
        "all that country's, and saying so in configuration costs nothing. Geography is a first-class "
        "scoring dimension, and leaving it to be inferred is why roughly a third of an earlier corpus "
        "carried none at all.",
    ])

    d.h2("3.3 What comes free, and the two things that do not")
    d.p(
        "Free, and this is the point of the seam: content hashing and deduplication, syndication "
        "collapse, language detection, tier assignment, relevance gating against the closed "
        "vocabulary, six-way classification, theme clustering, evidence attachment, all five "
        "attractiveness components, momentum, and inclusion in every downstream artefact from the "
        "brief to the portfolio plan. **Nothing in the pipeline needs to know that a new source "
        "exists.**"
    )
    d.p("Two things are not free, and both should be planned for rather than discovered:", after=5)
    d.callout(
        "1 · Scores move for every space, not only the ones the new source touches",
        [
            "Market signal strength is relative: a space's trailing-window signal count is divided by "
            "the largest such count in the live corpus. Add a large source and that denominator rises, "
            "so every existing space's component falls — correctly, but visibly.",
            "The consequence is procedural, not technical. A source change should be run as a "
            "**re-scoring at a fixed reference date**, so the before and the after are comparable, "
            "rather than dropped into a normal refresh where a source change and a genuine market "
            "movement are indistinguishable in the result.",
        ],
        fill=SH_BLUE, accent=BLUE,
    )
    d.callout(
        "2 · A new source has no history, and momentum needs six periods",
        [
            f"The replay archive ({n('raw_items')} raw items) makes a past date a re-run rather than a "
            "re-fetch — but only for material already collected. Backfill depth on a newly added "
            "source is bounded by whatever its API exposes, and for several of the news sources that "
            "is weeks, not years.",
            "Momentum is a least-squares slope over six trailing periods. A source added today "
            "contributes to it honestly only once six periods have passed; before that it inflates "
            "the level without informing the trend. Adding a source shortly before a decision "
            "milestone is therefore worse than adding it well before, or well after.",
        ],
        fill=SH_BLUE, accent=BLUE,
    )

    d.h2("3.4 The sources that are catalogued and dark")
    d.p(
        f"{S['dark']} of the {S['total']} catalogued sources are not enabled. The catalogue is the "
        "requirements record rather than only the runtime wiring, so each is listed with the reason it "
        "is dark. What is useful about the list is that the blockers are of three different kinds, and "
        "only one of them is engineering."
    )
    d.table(
        ["Source", "Blocked on", "Kind", "What enabling it buys"],
        [
            ["Adzuna job postings", "`ADZUNA_APP_ID` / `ADZUNA_APP_KEY`", "**Credentials**",
             "An enterprise hiring an OT-security engineer has committed budget months before a tender "
             "exists. Dated, employer-named, geo-tagged — and almost nobody uses postings for radar "
             "work. The connector is written and runs the moment the keys exist."],
            ["GitHub repository activity", "`GITHUB_TOKEN`", "**Credentials**",
             "Research sources measure what is being *studied*; repository activity measures what is "
             "being *built*, a later point on the same maturity curve. Written, tested, waiting."],
            ["Patent databases", "EPO OPS registration or BigQuery credentials", "**Procurement**",
             "The only observational input to technology ownership, which today is a portfolio-level "
             "flag from configuration — true for most spaces, and therefore a statement about Orange's "
             "portfolio in general rather than about that space. The CPC subsets the join needs are "
             "already in the technology vocabulary."],
            ["Standards bodies (3GPP, ETSI, O-RAN, GSMA)",
             "None of them publishes a feed; reaching them means parsing work plans and programmes",
             "**Engineering + terms**",
             "Standardisation activity is the earliest credible signal for network technology. This is "
             "a scraper with terms to clear, not a connector — the honest reason it is dark."],
            ["ENISA · EC foresight", "Terms of use, and an empty feed list", "**Terms**",
             "Foresight material maps cleanly onto the *Later* horizon, which is currently the thinnest "
             "of the three."],
            ["Eurostat ICT usage · Structural Business Statistics",
             "Nothing — they are loaded on the reference-data path, not as signals",
             "**Bookkeeping**",
             f"Already contributing: {n('observations')} observations behind every bottom-up market "
             "size. The catalogue rows should be reconciled so coverage reporting stops implying "
             "sizing has no inputs."],
            ["Internal signal intake",
             "Not a connector problem — records arrive by `radar internal add`, are inert until "
             "moderated, and reach the store on promotion",
             "**Organisational**",
             f"The one evidence class that says what Orange's own people are hearing, and the only one "
             f"no competitor can copy. {n('internal')} records exist. See 7I."],
        ],
        widths=[3.6, 4.0, 2.3, 6.7], size=8.5,
    )
    d.p(
        "Three sets of credentials and one registration would close five of these, and none of the "
        "four is an engineering task. That is the useful thing the table says.",
        after=10,
    )

    # ----------------------------------------------------------------- 4
    d.h1("4 · Connecting CRM")
    d.p(
        "CRM integration was deferred by the briefing, and the stated reason was that public assets "
        "give a sufficient right-to-win proxy. That was the right call for an MVP and it is the wrong "
        "call for a second version, because the proxy is now the binding constraint on the half of the "
        "radar that says whether Orange can win. This section says what CRM changes, how it would be "
        "connected in three different postures of data governance, and — the part that is usually "
        "underestimated — where the actual work is."
    )

    d.h2("4.1 What CRM actually changes, component by component")
    d.p(
        "Right to win is a deterministic structured lookup over the business graph — seven components, "
        "no model, every weight from configuration with a named owner. CRM does not change the "
        "formula. It changes what two of the heaviest components are looking at, and it makes a third "
        "quantity elsewhere in the system stop being a proxy.",
        after=5,
    )
    d.table(
        ["Component / quantity", "Weight", "What it reads today", "What CRM changes"],
        [
            ["**Offer match**", "0.25",
             f"Whether a link to an Orange offer exists, and whether any of them is an exact fit. "
             f"{n('links_l0')} exact-fit links exist across {n('links')} in total.",
             "The offer catalogue does not change — but which offers actually *sell*, into which "
             "vertical, does. A catalogue claim becomes an observation."],
            ["**Reference density**", "0.20",
             "The **published** customer-story distribution — 94 stories across twelve industry "
             "labels, of which eighteen are individually named — crosswalked onto fifteen verticals.",
             "**The single largest accuracy gain available anywhere in the scoring model.** Orange's "
             "real reference base is larger than its published one by orders of magnitude. The formula "
             "stays; the denominator becomes real; the evidence-gap warning starts firing on genuine "
             "gaps rather than on publication policy."],
            ["Partner coverage", "0.15", "Partner tier as an edge property, from configuration",
             "Little directly — but joint-deal history would let tier be *observed* rather than "
             "declared, and tiers change more often than configuration files do."],
            ["Compliance fit · Capability depth · External validation · Technology ownership",
             "0.12 · 0.12 · 0.08 · 0.08",
             "Certifications, capability-pool headcount, analyst positions, a portfolio-level flag",
             "Nothing. These are correctly sourced from elsewhere, and CRM should not be allowed near "
             "them."],
            ["**Obtainable share (SOM)**", "—",
             "Contract values observed from public procurement, because that is the only "
             "*attributable* source of contract values that exists.",
             "Orange's own won-deal distribution replaces a public-sector proxy with private-sector "
             "evidence. Every size in the radar currently rests on the assumption that the two "
             "resemble each other."],
            ["**Plan sequencing**", "—",
             "Entry year is bounded by entry slots and capability headcount",
             "Median cycle length by vertical turns a sequence into a schedule — the difference "
             "between \"do this second\" and \"this lands in year three\"."],
            ["**Competitive intensity**", "—",
             "A weighted count over a named competitor list, banded",
             "Lost-deal reason codes make it an observation of who actually turns up in a deal, which "
             "is a different and better quantity than who publishes a web page."],
        ],
        widths=[4.0, 1.7, 5.0, 5.9], size=8.5,
    )

    d.h2("4.2 Three integration patterns")
    d.p(
        "These are ordered by what Orange can approve fastest, which is not the same as ordered by "
        "value — although in this case the fastest to approve is also where most of the value is.",
        after=5,
    )

    d.h3("Pattern A · Aggregate export — recommended first")
    d.p(
        "A scheduled export of counts and sums, with no record-level data at all. One row per "
        "(industry × product family × country × year × stage), carrying opportunity count, won count, "
        "contract value in bands, median cycle days and win rate. No account names, no contact names, "
        "no per-deal amounts."
    )
    d.p(
        "It lands in `reference_observations`, beside the Eurostat observations, with a "
        "`reference_series` row recording the extract, its owner and its as-of date. Everything "
        "downstream then works unchanged — basis grading, confidence taken as the worst factor rather "
        "than an average, the explain modal, the version stamp — because that table already models "
        "exactly this shape: a value with coordinates and a provenance. The privacy argument is short, "
        "which is the point: there is no personal data in a count."
    )
    d.p(
        "Effort is a crosswalk, a loader and an agreed export from the CRM team. Two to three weeks of "
        "engineering once the export is agreed; the agreement is the longer pole."
    )

    d.h3("Pattern B · Named-account graph edges")
    d.p(
        "A won deal becomes a `reference` node; an active opportunity becomes an `opportunity` node; "
        "each links to the spaces it belongs to as a typed `opportunity_link` carrying evidence and a "
        "confidence, exactly like every other link in the graph. This is what makes a brief say "
        "*Orange has done this, here, for this customer* rather than *Orange has offers in this area* "
        "— and it is what a salesperson actually needs in the room."
    )
    d.p(
        "It is also the pattern that puts customer names into the store. That needs a data-protection "
        "assessment and a retention rule **before** the schema change, not after. Everything held "
        "today is either public or a username and a password verifier; Pattern B changes the category "
        "of the system, and the right time to notice that is now."
    )

    d.h3("Pattern C · Live query at read time")
    d.p(
        "Ask the CRM API when the panel opens, store nothing. Best governance posture, and it must not "
        "be the primary path: a score that cannot be reproduced from stored inputs breaks the "
        "reproducibility guarantee the whole tool rests on — a reviewer outside the project has to be "
        "able to reconstruct why a topic holds its rank, and \"the CRM said so in August\" is not a "
        "reconstruction. Use it for a *who owns this account* panel next to the brief, where staleness "
        "matters more than reproducibility, and nowhere in scoring."
    )
    d.table(
        ["Pattern", "Gives", "Costs", "Breaks", "When"],
        [
            ["**A · Aggregate export**", "Real reference density, real win rates, real cycle times",
             "A crosswalk, a loader, an agreed export", "Nothing", "**Phase 1**"],
            ["**B · Named-account edges**", "Per-space proof points a salesperson can name",
             "A DPIA, a retention rule, node and link types",
             "The system's current no-personal-data posture — deliberately, with controls",
             "Phase 2, gated on the assessment"],
            ["**C · Live query**", "Always-current account ownership", "An integration and its uptime",
             "Reproducibility, if it is ever allowed to touch a score",
             "Optional, presentation only"],
        ],
        widths=[3.5, 4.0, 3.4, 4.0, 2.6], size=8.5,
    )

    d.h2("4.3 The crosswalk is the work")
    d.p(
        "CRM records are tagged by product code and by account industry classification. The radar is "
        "indexed by vertical × use case × technology. Nothing joins those two without a mapping, and "
        "the mapping is a judgement call that has to be owned by a person — which is why it is the "
        "work, and the connector is the easy half."
    )
    d.p(
        "The project already has the pattern for exactly this. Four crosswalk files exist — "
        "`cpv_to_vertical`, `cpv_to_use_case`, `vertical_to_nace`, `technology_to_adoption` — each "
        "versioned, each with a named owner, each auditable. Two more in the same shape are what CRM "
        "needs:",
        after=5,
    )
    d.bullets([
        "`config/crosswalks/crm_product_to_offer.csv` — product code to offer id, with a share where "
        "a code spans more than one offer.",
        "`config/crosswalks/crm_industry_to_vertical.csv` — the CRM's industry classification to the "
        "fifteen verticals, with a share where the labels do not nest.",
    ])
    d.callout(
        "The rule that makes a crosswalk safe",
        [
            "**An unmapped code is reported, never silently dropped.** A silently dropped product "
            "family looks exactly like a vertical where Orange does no business — and it will be "
            "read as evidence, because everything else in this tool is.",
            "The same rule already governs the CPV crosswalks and the competitor register. It is "
            "cheap to honour at build time and impossible to reconstruct afterwards.",
        ],
    )
    d.p(
        "On size: a product catalogue in the hundreds of codes and an industry classification of a "
        "dozen or so labels means a first pass covering the large majority of revenue is a day's work "
        "with the right person in the room. The tail is where the argument is, and the tail is also "
        "where the interesting spaces live — so it should be scheduled as a recurring review rather "
        "than a one-off, with the same curator who adjudicates links."
    )

    d.h2("4.4 The wall that has to stay up")
    d.callout(
        "CRM data belongs in right to win, and nowhere else",
        [
            "Feeding CRM data into **attractiveness** would make the radar rate a market as attractive "
            "because Orange already sells there. That is precisely the failure a radar exists to "
            "prevent: it turns a discovery tool into a mirror of the current business, and it does so "
            "invisibly, because the resulting rankings look reassuringly sensible.",
            "The two scores are kept separate for exactly this reason, and the CRM connection is the "
            "moment the temptation to combine them will be strongest — the request will arrive as "
            "\"can we just weight it a little by where we already have traction\".",
            "The answer is that traction already has a score, it is called right to win, and both are "
            "on the chart.",
        ],
        fill=SH_RED, accent=RED,
    )
    d.p(
        "A second, quieter version of the same mistake is worth naming: CRM history is backward "
        "looking. It is the correct input to *can we win here*, and it is a systematically misleading "
        "input to *is this worth entering*. Where the two disagree — a space with strong external "
        "momentum and no Orange history — the disagreement is the finding, not an error to be "
        "smoothed away.",
        after=10,
    )

    # ----------------------------------------------------------------- 5
    d.h1("5 · Closing the rest of the not-built list")
    d.p(
        "The slide names six exclusions. CRM is section 4; the other five are here, each with the "
        "smallest version of the work that would actually change something, because the full version "
        "of each is a research programme and the smallest version is usually a fortnight."
    )

    d.h2("5.1 Calibration and backtesting — the highest-value modelling work")
    d.p(
        "Nothing in the radar has been calibrated and nothing has been backtested. The weights are the "
        "briefing's indicative figures and no outcome data has ever moved them. The replay harness a "
        "backtest needs is built and works — replay at a past reference date, publication-date leakage "
        "control, retained raw archives — and the evaluation metrics are not implemented, so the "
        "radar cannot yet answer the question that would justify its own ordering."
    )
    d.p("The smallest useful version, and it is genuinely small:", after=5)
    d.numbers([
        "Replay at three past reference dates spread across the archive.",
        "For each, measure **rank stability** of the top quartile — how much does the ordering move "
        "when the corpus grows by a known amount?",
        "For each, measure **precision** against what has since happened: of the spaces that scored "
        "high then, how many have since been tendered, funded, regulated into existence, or entered "
        "by a named competitor?",
    ])
    d.p(
        "That is one engineer for about two weeks, and it converts *a defensible, transparent, "
        "asserted ordering* into *a measured one*. Only after that should weights move — and any move "
        "requires a new weight-set id, because scores computed under different weights are not "
        "comparable and the interface must keep refusing to plot them together."
    )

    d.h2("5.2 Learned per-role ranking")
    d.p(
        f"This needs 300–600 expert pairwise comparisons. The capture widget ships and the feedback "
        f"table holds {n('feedback')} row, which tells you that the elicitation is the work and the "
        "model is not. Eight to ten people, two weeks, a single repeated question in role context — "
        "*which of these two would you work on first?* — produces the labels. The resulting ranker "
        "must remain an **ordering over the same published components**, never a new opaque score, or "
        "it takes the explain modal down with it."
    )

    d.h2("5.3 Patents, per-role authorisation, rate limiting, ROI")
    d.table(
        ["Item", "What it actually is", "Size", "Why it matters"],
        [
            ["**Patent connector**", "An EPO OPS registration or BigQuery credentials — a procurement "
             "action. The CPC subsets are already carried in the technology vocabulary and the mapping "
             "is designed.", "Days, once registered",
             "Turns technology ownership from a portfolio-level flag into an observation about *this* "
             "space. It is 8% of right to win and currently carries no information at the space level."],
            ["**Per-role authorisation**",
             "Sign-in answers *who*; it does not answer *may they*. Every signed-in account can move a "
             "stage, delete a space and spend model budget. The role vocabulary already exists — "
             "strategist, sales, presales — so this is a role on the user record and a check at the "
             "API boundary.", "Days",
             "**This is the item that decides whether the radar can be deployed beyond a pilot.** It "
             "is small, unglamorous, and blocking."],
            ["**Rate limiting on generation**",
             "A budget guard on the endpoints that spend model money, at the same boundary as the "
             "authorisation check.", "An afternoon",
             "Sign-in bounds *who* can reach those endpoints, not *how often*. One enthusiastic "
             "afternoon of clicking is currently unbounded spend."],
            ["**ROI on a plan**",
             "Still not honest, and should stay unbuilt. There is no cost data at the granularity a "
             "space needs, anywhere the pipeline can reach.", "Blocked on data, not effort",
             "The defensible substitute is **cost bands per capability pool** with a named finance "
             "owner, stamped with an economics version — which yields a payback *range*. A range that "
             "says so is worth more than a ratio that invented its denominator."],
        ],
        widths=[3.4, 6.4, 2.4, 5.2], size=8.5,
    )

    # ----------------------------------------------------------------- 6
    d.h1("6 · The decisions Orange owes the project")
    d.p(
        "These are not engineering items and they cannot be closed by the delivery team. Each is "
        "listed with what it blocks and what it costs to leave open, because the cost of an open "
        "decision here is not delay — it is that the radar keeps producing output that is one "
        "assumption away from being wrong, and says so in the interface every time."
    )
    d.table(
        ["Decision", "What is blocked", "Cost of leaving it open"],
        [
            ["**Who is the curator?**",
             f"{n('links')} links are machine-proposed. **{n('links_confirmed')} are confirmed and "
             f"{n('links_rejected')} rejected**; {n('pattern_decisions')} link patterns have been "
             "adjudicated. Right to win is a structured lookup over exactly those links.",
             "Every right-to-win score in the radar rests on evidence nobody has signed. The code to "
             "record a curator's decision exists and is tested. The curator does not."],
            ["**Margin by portfolio distance**",
             "The Planner's five-year profit projection. One table from Orange finance replaces the "
             "planning bands anchored on the filed segment margin.",
             "Varying margin by distance rather than holding it flat moves five-year profit by about "
             "**1.66×**, and revenue concentrates at the nearest distance band — so this single table "
             "dominates the answer more than build cost does."],
            ["**Headcount free for new work**",
             "The Planner's binding constraint in most runs.",
             "The capability pools describe who *exists*, not who is *free*. The constraint that binds "
             "first in most plans is therefore currently a guess, and the plan is bounded by it."],
            ["**Terms of use**",
             f"{S['pending']} of the {S['enabled']} enabled sources carry `terms_checked: pending`.",
             "A Sprint 0 blocker for anything beyond a prototype, and **the cheapest item on this "
             "list to close.** Storage is by reference with bounded extracts, so the answer is likely "
             "to be yes — but it has to be asked."],
            ["**Refresh cadence**",
             "Connector design, cost, and how momentum is measured.",
             "Drives cost more than any other single choice, and is currently fourteen days by "
             "default rather than by decision. See 8.4 — the price of each option is not yet "
             "instrumented."],
            ["**Sovereign deployment**",
             "Whether an external model API may be used at all.",
             "The client abstraction supports a local model today and that path has not been exercised "
             "at corpus scale. Answering this late is a materially different job from answering it "
             "now."],
            ["**A browser user agent for competitor profiling**",
             f"{F['competitor_unprofiled']} of {n('competitor_profiles')} competitors are unprofiled; "
             "six sites answer 403 to a declared automated client.",
             "Recorded as a refusal rather than routed around, which is the correct default. The "
             "thinning falls hardest on security spaces. The decision is Orange's to make, not the "
             "pipeline's to work around."],
            ["**Is four years the right contract duration?**",
             "Every market size computed from procurement observations.",
             "Published contract values are whole-contract; annualising them needs a duration. Every "
             "size in the radar moves inversely with this number, and an Orange bid team has a better "
             "one than the printed assumption."],
        ],
        widths=[3.6, 5.8, 7.0], size=8.5,
    )
    d.callout(
        "The pattern in this table",
        "Six of the eight are answered by a person accepting ownership of something — a curator, a "
        "finance table, a headcount, a legal review. None needs a sprint. Together they move the "
        "radar from *a working system with placeholders* to *a system Orange stands behind*, and no "
        "amount of engineering substitutes for them.",
    )

    # ----------------------------------------------------------------- 7
    d.h1("7 · New directions worth building")
    d.p(
        "Nothing in this section was asked for. Each is here because the machine that now exists makes "
        "it unusually cheap, and because it addresses something the current design does not: adoption, "
        "timing, or the sales-side question the radar currently answers backwards."
    )

    d.h3("7A · The refresh diff — turn the radar from a screen into a subscription")
    d.p(
        f"{n('refreshes')} refreshes are recorded, each with per-stage statistics, and every score row "
        "carries the refresh that produced it. What does not exist is the **delta**: which spaces "
        "changed band since last time, and which specific signal moved them. That is a query and a "
        "template — no model, no new data, no new seam."
    )
    d.p(
        "It is the highest ratio of value to effort in this document, and it fixes the problem that "
        "quietly kills tools like this one: **a tool people have to remember to open is a tool that "
        "gets opened twice.** A fortnightly note that says *six spaces moved, here is the notice that "
        "moved them* is read. A dashboard is not."
    )

    d.h3("7B · Invert the question — which of my accounts sits in a moving space")
    d.p(
        "The radar answers *which spaces are attractive*. A salesperson asks the transpose: *which of "
        "my accounts is in one*. With Pattern A's aggregate export and the industry crosswalk this is "
        "a join, not a project — and it converts a quarterly strategy artefact into a weekly sales "
        "one. This is the payoff for section 4 and the strongest argument for doing section 4 first."
    )

    d.h3("7C · Ask the evidence")
    d.p(
        f"A retrieval interface over the {n('signals')} stored signals, answering in citations. This "
        "is one of the few places where a model should read the corpus at query time, and the "
        "guardrails already exist and are tested: evidence binding, entailment checking, and the "
        "no-numbers rule enforced in the client rather than in a prompt. Constrained to answer only "
        "from stored signals, it inherits the property that makes the rest of the tool credible. The "
        "tempting version — where it also speculates — should be refused, and refused loudly, because "
        "it is the version that will be requested."
    )

    d.h3("7D · The regulatory deadline calendar")
    d.p(
        "Roughly a quarter of the catalogued sources are regulators, and regulation is the one signal "
        "class carrying a **dated future obligation**. NIS2, DORA, the Cyber Resilience Act, the AI "
        "Act and the data-act family all have compliance dates attached. Extract that date as a typed "
        "attribute on the signal — extraction over material already collected — and the radar can rank "
        "by *what becomes mandatory in the next eighteen months*."
    )
    d.p(
        "That is a question every enterprise customer is already asking, that no ranked innovation "
        "list answers, and that maps directly onto a sales conversation with its own urgency built in. "
        "It is the most differentiated idea in this document."
    )

    d.h3("7E · A partner-facing view")
    d.p(
        "Partner tier is already an edge property in the graph. A partner-scoped read — *where do "
        "Orange and this partner have joint right to win* — is a filter over a graph that exists, and "
        "it produces a co-sell artefact of exactly the kind partners routinely fund. Close to free, "
        "and it creates an external audience for the radar, which is the most reliable way to keep an "
        "internal tool maintained."
    )

    d.h3("7F · A competitor change feed")
    d.p(
        f"{n('competitor_pages')} competitor pages are stored with bounded extracts. Crawled once, "
        "they are a snapshot. Crawled on a cadence and diffed, they become a change feed — *this "
        "competitor added OT security to its manufacturing page in March* is a sharper and more "
        "actionable signal than the page itself, and the baseline is already in the store. The "
        "incremental cost is one more crawl and a comparison."
    )

    d.h3("7G · Close the loop back into the CRM")
    d.p(
        "Twelve pre-sales pieces per space already build in three formats each. Pushing the pack back "
        "as an attachment on the opportunity record means nobody has to leave the CRM to get it — and "
        "it produces the usage telemetry the learned ranking of 5.2 needs, including the question "
        "nobody can currently answer: which of the twelve pieces does anyone actually open?"
    )

    d.h3("7H · Per-country instances")
    d.p(
        "Geography is a first-class dimension and country managers will want their own corpus, their "
        "own plan and their own stage gate. One database file per tenant matches the storage model "
        "already in use, where a replay is a file copy. The cost is operational — deployment, backup, "
        "config drift between tenants — not architectural."
    )

    d.h3("7I · Internal signal intake, properly routed")
    d.p(
        f"{n('internal')} records exist in the one evidence class that says what Orange's own people "
        "are hearing. The mechanism is built — add, moderate, promote, tier 3 on purpose, because a "
        "conversation is not a published authoritative record however well informed. What is missing "
        "is a **route into it that fits how a salesperson actually works**: an email address, a Teams "
        "action, a one-field form on the brief itself."
    )
    d.p(
        "This is the cheapest way to make the radar feel like it belongs to the people using it rather "
        "than to the team that built it — and it is the only evidence class no competitor can "
        "replicate, because it is the only one that is not public."
    )

    d.table(
        ["#", "Direction", "Depends on", "Effort", "What it changes"],
        [
            ["7A", "**Refresh diff / subscription**", "Nothing", "**Small**",
             "Adoption. The difference between a tool that is opened and one that is remembered."],
            ["7B", "**Account matching**", "§4 Pattern A", "Small once A lands",
             "Turns a quarterly artefact into a weekly one."],
            ["7C", "Ask the evidence", "Nothing", "Medium", "Makes the corpus usable by people who "
             "will never learn the taxonomy."],
            ["7D", "**Regulatory deadline calendar**", "Extraction over existing signals", "Medium",
             "Adds *when* to a system that currently answers *what* — and it is the most "
             "differentiated idea here."],
            ["7E", "Partner view", "Nothing", "**Small**", "An external audience, and a co-sell "
             "artefact partners fund."],
            ["7F", "Competitor change feed", "A recurring crawl", "Small", "Turns a snapshot into a "
             "signal."],
            ["7G", "Collateral into the CRM", "§4 Pattern A or C", "Medium",
             "Closes the loop, and produces the usage labels 5.2 needs."],
            ["7H", "Per-country instances", "Deployment work", "Medium", "Makes the radar a product "
             "rather than an instance."],
            ["7I", "**Internal intake route**", "Nothing", "**Small**",
             "The only evidence class no competitor can copy."],
        ],
        widths=[1.0, 3.6, 3.0, 2.2, 6.6], size=8.5,
    )

    # ----------------------------------------------------------------- 8
    d.h1("8 · Platform, operations and technical debt")

    d.h3("8.1 Storage — defend the choice, and name the trigger to change it")
    d.p(
        "A single-file relational store is the right size of solution here and should be defended "
        "rather than apologised for: the graph is thousands of nodes rather than millions, the serving "
        "profile is read-mostly with a single writer, discovery is a scheduled batch job, and a "
        "historical replay is a file copy rather than a restore procedure."
    )
    d.p(
        "What matters is naming the **trigger** to move, and it is not row count — it is concurrent "
        "writers. Per-country instances (7H) fit the current model exactly. Many people assessing and "
        "moving stages simultaneously inside one instance does not. A move would be a schema port; the "
        "two identity rules that make refreshes stable would carry unchanged."
    )
    d.p(
        "One deployment constraint belongs in the runbook rather than in folklore: the hosted instance "
        "runs in rollback-journal rather than write-ahead-log mode, because the application's file "
        "share cannot provide the shared memory WAL requires. That is a constraint on any future "
        "hosting choice, not a fix that happened once."
    )

    d.h3("8.2 Detect the silent zero")
    d.p(
        "Per-source errors are already recorded per refresh. What is **not** detected is a source that "
        "returns a perfectly healthy response and no items — the failure mode this corpus has already "
        "demonstrated once, and the one that degrades a radar without ever raising an error. A "
        "per-source expected-yield band, with an alert when a run falls outside it, is about a day's "
        "work and prevents the worst class of quiet decay."
    )

    d.h3("8.3 Instrument the cost of a refresh")
    d.p(
        "Model spend should be a stored number on each refresh row. Refresh cadence is one of the six "
        "open decisions, and it is currently being decided without a price attached to the "
        "alternatives — which is the one piece of information that would actually settle it."
    )

    d.h3("8.4 Exercise the sovereign path before it becomes a commitment")
    d.p(
        "The model client is provider-neutral and a local model is supported, but that path has not "
        "been run at corpus scale. The failure modes of a smaller local model appear first in the "
        "cheap, high-volume stages — classification and entailment — not in the prose everyone would "
        "think to inspect. Test it while it is still an option rather than after it has become a "
        "requirement."
    )

    d.h3("8.5 The documentation build is broken at the source")
    d.p(
        "`docs/generators/build_docs.py` still assembles the Functional Design Document, the Technical "
        "Architecture and the Speaker Notes from content modules in `docs/_build/`, and those modules "
        "are no longer in the tree. The `.docx` outputs are current and correct; they simply cannot be "
        "rebuilt from source as things stand. Either restore the content modules or re-point the "
        "generator — and do it before the next edit is needed under time pressure, which is the "
        "circumstance in which this will otherwise be discovered."
    )

    d.h3("8.6 Personal data is a category change, not a volume change")
    d.p(
        "Everything in the store today is either public or a username and a verifier, which is why the "
        "data-protection position is currently short enough to fit in a paragraph. Section 4 Pattern B "
        "changes that. The assessment and the retention rule belong before the schema change, and the "
        "cheapest possible version of this document's advice is: **do Pattern A first, and find out "
        "how much of the value arrives without ever storing a customer name.**"
    )

    # ----------------------------------------------------------------- 9
    d.h1("9 · A suggested sequence")
    d.p(
        "Four horizons. The first contains no features at all, which is deliberate: it is the set of "
        "things that decide whether the radar can be deployed and trusted, and each one is small."
    )
    d.table(
        ["Horizon", "Work", "Unlocks", "Depends on"],
        [
            ["**Sprint 0**\n(weeks)",
             "Clear terms of use on the enabled sources · name the curator · per-role authorisation · "
             "rate limiting on generation · three sets of source credentials · reconcile the "
             "reference-data catalogue rows",
             "Deployment beyond a pilot; signed evidence behind right to win; two more live sources",
             "Orange decisions, not engineering capacity"],
            ["**Phase 1**\n(one quarter)",
             "CRM Pattern A and the two crosswalks · the refresh diff and subscription (7A) · backtest "
             "metrics at three dates (5.1) · patent registration started · silent-zero alerting (8.2) "
             "· internal intake route (7I)",
             "Real reference density and real win rates; a measured ordering rather than an asserted "
             "one; a tool people receive rather than visit",
             "Sprint 0 · an agreed CRM export · a named crosswalk owner"],
            ["**Phase 2**",
             "Account matching (7B) · regulatory deadline calendar (7D) · named-account edges "
             "(Pattern B, post-assessment) · partner view (7E) · competitor change feed (7F) · "
             "pairwise elicitation begins (5.2)",
             "The sales-side inversion; timing as a first-class dimension; per-space proof points a "
             "salesperson can name",
             "Phase 1 · a completed data-protection assessment"],
            ["**Phase 3**",
             "Calibrated weights under a new weight set · learned per-role ranking · per-country "
             "instances (7H) · collateral into the CRM (7G)",
             "Scores moved by measured outcomes rather than by a briefing's indicative figures",
             "Phase 2 · 300–600 captured comparisons · measured backtest results"],
        ],
        widths=[2.4, 5.6, 4.6, 4.0], size=8.5,
    )
    d.callout(
        "What the ordering is really saying",
        [
            "The two items that most change what the radar is worth are **a named curator** and **a "
            "CRM aggregate export**. Neither is an engineering task. Both are somebody agreeing to own "
            "something.",
            "Everything in Phase 3 is more interesting to build and worth less until the two above "
            "have happened — learned weights over unsigned evidence would be a more sophisticated way "
            "of being confidently wrong.",
        ],
    )

    # ----------------------------------------------------------------- 10
    d.h1("10 · What should stay unbuilt")
    d.p(
        "A roadmap that only adds is a roadmap that erodes the thing being extended. These six hold "
        "regardless of what else on this list is built, and each one will come under pressure from "
        "something proposed above.",
        after=5,
    )
    d.bullets([
        "**No model-generated numbers, ever** — including on every new surface added above. The rule "
        "is enforced in the client, backstopped by a check over generated text, and tested. A new "
        "feature that needs an exception does not need an exception; it needs arithmetic.",
        "**CRM into right to win only, never into attractiveness.** The request will arrive as \"weight "
        "it a little by where we already have traction\". Traction already has a score.",
        "**No ROI until cost data exists** at the granularity a space needs. A payback band with a "
        "named owner, not a ratio with an invented denominator.",
        "**No headless browser to reach a handful of competitor sites.** The refusal is the record, "
        "and a recorded refusal is more useful than a quietly circumvented one — it is a decision "
        "Orange can take, which is the point of writing it down.",
        "**No comparison across a weight-set, sizing-version or economics-version boundary** — "
        "including the comparisons the refresh diff (7A) will make extremely tempting, since a diff "
        "across a re-weighting looks exactly like a market movement.",
        "**No learned scoring before labels exist.** The transparent baseline is not a placeholder to "
        "be replaced quietly; it is what makes the explanation possible, and a learned model that "
        "cannot be explained is a worse product even when it is more accurate.",
    ])
    d.p("", after=4)
    d.rule(ORANGE)
    d.p(
        "The radar's distinguishing property is not that it ranks well. It is that every number "
        "decomposes into named inputs, every claim resolves to a document somebody published on a "
        "date, and where the evidence runs out it publishes nothing rather than an estimate. Each step "
        "in this document is scoped to keep that true — which is why the shortest section is the one "
        "about what not to build, and why it is the last word.",
        italic=True, color=GREY, before=8,
    )

    d.save(OUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
